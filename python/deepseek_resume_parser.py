import sys
import json
import re
import os
import requests
from PyPDF2 import PdfReader
import traceback
import time  # Import for retry mechanism

# Try to import docx library for handling Word documents
try:
    import docx
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX parsing will be limited.", file=sys.stderr)

# Sample structure of the desired output JSON
SAMPLE_OUTPUT_STRUCTURE = """
{
  "name": "John Doe",
  "contact_info": {
    "email": "john.doe@example.com",
    "phone": "+1 (123) 456-7890",
    "linkedin": "https://linkedin.com/in/johndoe",
    "github": "https://github.com/johndoe"
  },
  "education": [
    {
      "institution": "Stanford University",
      "degree": "MS Computer Science",
      "date": "2016-2018",
      "description": "MS Computer Science, Stanford University, 2016-2018"
    }
  ],
  "experience": [
    {
      "company": "Google",
      "position": "Software Engineer",
      "date": "2018-Present",
      "description": "Software Engineer at Google, 2018-Present"
    }
  ],
  "skill": {
    "technical_skills": ["Python", "JavaScript", "Machine Learning"],
    "soft_skills": [],
    "tools": []
  }
}
"""

def extract_pdf_text(pdf_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            print(f"Warning: Extracted PDF text is empty or contains only whitespace", file=sys.stderr)
            # Try alternate method if the primary method returns empty text
            return extract_pdf_text_fallback(pdf_path)
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        # Try fallback method
        return extract_pdf_text_fallback(pdf_path)

def extract_pdf_text_fallback(pdf_path):
    """Fallback method for extracting text from PDF files"""
    try:
        print(f"Trying fallback method for PDF: {pdf_path}", file=sys.stderr)
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            
            # Try a different approach with extracting text
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                
                # Extract text with different settings
                # Some PDFs work better with this method
                page_text = ""
                try:
                    # Try standard extraction first
                    page_text = page.extract_text()
                except:
                    # If that fails, try raw extraction
                    if hasattr(page, '_extract_text'):
                        page_text = page._extract_text()
                
                if page_text:
                    text += page_text + "\n"
                    
        return text
    except Exception as e:
        print(f"Fallback PDF extraction failed: {e}", file=sys.stderr)
        return ""

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        if DOCX_AVAILABLE:
            print(f"Using python-docx to extract text from {docx_path}", file=sys.stderr)
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            full_text = []
            
            # Get all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            # Create well-formatted text with proper spacing
            text = '\n'.join(full_text)
            
            # Clean up excessive whitespace
            text = re.sub(r'\n\s*\n', '\n\n', text)
            return text
        else:
            # Fallback: Try to read with binary mode and decode
            print(f"Fallback method for DOCX: reading as binary", file=sys.stderr)
            with open(docx_path, 'rb') as f:
                content = f.read()
                # Try to extract plain text portions
                text_parts = []
                current_text = ""
                for byte in content:
                    if 32 <= byte <= 126:  # ASCII printable characters
                        current_text += chr(byte)
                    elif byte in (10, 13):  # newline, carriage return
                        current_text += '\n'
                    elif current_text:
                        text_parts.append(current_text)
                        current_text = ""
                
                if current_text:
                    text_parts.append(current_text)
                
                # Filter out very short parts and join
                text = '\n'.join([part for part in text_parts if len(part) > 5])
                return text
    except Exception as e:
        print(f"Error reading DOCX: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        return ""

def clean_text(text):
    """Clean extracted text by removing special characters and normalizing whitespace"""
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n+', '\n', text)
    # Normalize whitespace within lines (not removing newlines)
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove non-printable characters except for newlines
    text = ''.join(c for c in text if c == '\n' or (32 <= ord(c) <= 126))
    # Cut off text if it's extremely long to stay within model context limits
    max_length = 12000  # Conservative maximum
    if len(text) > max_length:
        print(f"Warning: Text too long ({len(text)} chars), truncating to {max_length} chars", file=sys.stderr)
        return text[:max_length]
    return text

def send_to_deepseek_ai(text, max_retries=3):
    """Send the extracted text to DeepSeek AI via OpenRouter to parse the resume and return JSON"""
    print("Sending extracted text to DeepSeek AI for parsing", file=sys.stderr)
    
    # Use hardcoded OpenRouter API token
    openrouter_token = "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"
    
    print("Using deepseek with openrouter", file=sys.stderr)
    
    # Define the prompt for OpenRouter
    prompt = f"""
    You are an expert resume parser. Extract key information from the following resume text 
    and format it according to the EXACT specified JSON structure. Be precise and accurate.

    *** CRITICALLY IMPORTANT: You MUST follow these rules ***
    1. Output ONLY valid JSON matching the exact structure below - no other text
    2. The skills section MUST be under the key "skill" (NOT "skills")
    3. DO NOT add any additional top-level fields that are not in the sample structure
    4. DO NOT include technical_skills, soft_skills, tools, email, phone, linkedin, or github as top-level fields
    5. DO NOT duplicate data between nested objects and the top level
    6. All data for a given type MUST be in its designated field only (e.g., contact info in contact_info object only)
    7. Follow the EXACT structure shown below - do not deviate or add extra fields

    The output MUST EXACTLY match this structure:
    {SAMPLE_OUTPUT_STRUCTURE}

    Here is the resume text to parse:
    
    {text}
    
    IMPORTANT INSTRUCTIONS FOR SKILL EXTRACTION:
    1. Be exhaustive in extracting technical skills, tools, languages, and soft skills
    2. Categorize skills properly into technical_skills, soft_skills, and tools categories
    3. Include ALL programming languages, frameworks, platforms, and technologies mentioned
    4. Look for skill sections, but also extract skills mentioned throughout the entire resume
    5. For technical roles (DevOps, Cloud, etc.), include relevant cloud platforms, CI/CD tools, containers, etc.
    6. For Oracle/BI roles, include all Oracle-related technologies and reporting tools
    7. The skills section is critical - DO NOT leave this section empty or incomplete
    8. The skills key in the JSON MUST be "skill" (not "skills")

    Format your response ONLY as the exact JSON structure shown above - no explanations, no markdown, just the JSON.
    """
    
    # OpenRouter API endpoint
    url = "https://openrouter.ai/api/v1/chat/completions"
    
    # Request headers
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openrouter_token}",
        "HTTP-Referer": "https://talnurt.com"  # Optional: helps with OpenRouter analytics
    }
    
    # Request payload
    payload = {
        "model": "deepseek/deepseek-r1:free",
        "messages": [
            {"role": "system", "content": "You are an expert resume parser that extracts structured information into JSON."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.1,
        "max_tokens": 3000
    }
    
    # Initialize variables for retry logic
    retry_count = 0
    wait_time = 2  # Start with 2 seconds wait
    
    while retry_count < max_retries:
        try:
            # Send request to the API
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            
            if response.status_code == 200:
                data = response.json()
                
                # Debug logging
                print("\n--- DEBUG: API Response Status ---", file=sys.stderr)
                print(f"Status code: {response.status_code}", file=sys.stderr)
                
                if 'choices' in data and len(data['choices']) > 0 and 'message' in data['choices'][0]:
                    ai_response = data["choices"][0]["message"]["content"].strip()
                    
                    # Debugging: Print the first 100 characters of the response
                    print(f"Response excerpt: '{ai_response[:100]}...'", file=sys.stderr)
                    
                    # If the response is empty, try another model
                    if not ai_response:
                        print("Received empty response from DeepSeek. Trying fallback method.", file=sys.stderr)
                        return generate_fallback_json(text)
                    
                    # Extract JSON from the AI response
                    json_match = re.search(r'```json\s*([\s\S]*?)\s*```', ai_response)
                    if json_match:
                        json_text = json_match.group(1).strip()
                        print(f"Found JSON in code block: {json_text[:50]}...", file=sys.stderr)
                    else:
                        # If no code blocks, try to extract JSON directly
                        # First check if the response starts with a curly brace
                        if ai_response.startswith('{') and '}' in ai_response:
                            json_text = ai_response
                            print("Response appears to be direct JSON", file=sys.stderr)
                        else:
                            # Look for JSON-like structure in response
                            json_match = re.search(r'({[\s\S]*})', ai_response)
                            if json_match:
                                json_text = json_match.group(1).strip()
                                print(f"Extracted JSON using regex: {json_text[:50]}...", file=sys.stderr)
                            else:
                                print("Error: Could not extract JSON from response", file=sys.stderr)
                                print(f"Raw response: '{ai_response}'", file=sys.stderr)
                                # Try to validate if the entire response is valid JSON
                                try:
                                    json.loads(ai_response)
                                    json_text = ai_response
                                    print("The entire response seems to be valid JSON", file=sys.stderr)
                                except:
                                    print("Falling back to generate basic JSON from text", file=sys.stderr)
                                    return generate_fallback_json(text)
                    
                    # Try to validate the JSON before returning
                    try:
                        parsed_json = json.loads(json_text)
                        print(f"Successfully validated JSON with {len(json_text)} characters", file=sys.stderr)
                        return json_text
                    except json.JSONDecodeError as e:
                        print(f"Error: Invalid JSON: {e}", file=sys.stderr)
                        print(f"JSON text: '{json_text}'", file=sys.stderr)
                        # Try to clean the JSON string
                        cleaned_json = clean_json_string(json_text)
                        try:
                            json.loads(cleaned_json)
                            print("Successfully cleaned and validated JSON", file=sys.stderr)
                            return cleaned_json
                        except json.JSONDecodeError:
                            # Try to repair truncated JSON
                            repaired_json = repair_truncated_json(json_text)
                            try:
                                json.loads(repaired_json)
                                print("Successfully repaired truncated JSON response", file=sys.stderr)
                                return repaired_json
                            except json.JSONDecodeError:
                                print("Could not repair JSON. Using fallback method.", file=sys.stderr)
                                return generate_fallback_json(text)
                else:
                    print(f"Error: Unexpected API response format: {data}", file=sys.stderr)
                    return generate_fallback_json(text)
            elif response.status_code == 429:  # Rate limit error
                retry_count += 1
                print(f"Rate limit error. Retrying in {wait_time} seconds (attempt {retry_count}/{max_retries})", file=sys.stderr)
                time.sleep(wait_time)
                # Exponential backoff
                wait_time *= 2
            else:
                print(f"Error: API call failed with status {response.status_code}: {response.text}", file=sys.stderr)
                # For non-rate-limit errors, we'll still retry but not increase wait time
                retry_count += 1
                if retry_count < max_retries:
                    print(f"Retrying in {wait_time} seconds (attempt {retry_count}/{max_retries})", file=sys.stderr)
                    time.sleep(wait_time)
                else:
                    print("All retries failed. Using fallback method.", file=sys.stderr)
                    return generate_fallback_json(text)
        except requests.exceptions.RequestException as e:
            print(f"Network error: {e}", file=sys.stderr)
            retry_count += 1
            if retry_count < max_retries:
                print(f"Retrying in {wait_time} seconds (attempt {retry_count}/{max_retries})", file=sys.stderr)
                time.sleep(wait_time)
                wait_time *= 2
            else:
                print("All retries failed due to network issues. Using fallback method.", file=sys.stderr)
                return generate_fallback_json(text)
    
    # If we've exhausted all retries
    print("All API attempts failed. Using fallback method.", file=sys.stderr)
    return generate_fallback_json(text)

def clean_json_string(json_str):
    """Clean and repair malformed JSON strings"""
    # Remove any leading/trailing whitespace
    json_str = json_str.strip()
    
    # Replace curly quotes with straight quotes
    json_str = json_str.replace('"', '"').replace('"', '"')
    
    # Fix missing quotes around keys
    json_str = re.sub(r'(\{|\,)\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', json_str)
    
    # Fix single quotes to double quotes
    json_str = re.sub(r"'([^']*)'", r'"\1"', json_str)
    
    # Fix trailing commas in arrays/objects
    json_str = re.sub(r',\s*\}', '}', json_str)
    json_str = re.sub(r',\s*\]', ']', json_str)
    
    # Fix missing quotes around string values
    json_str = re.sub(r':\s*([a-zA-Z_][a-zA-Z0-9_\s]*)(,|})', r':"\1"\2', json_str)
    
    # Attempt to fix truncated JSON (ending with a comma or an unclosed quote)
    json_str = re.sub(r',$', '', json_str)
    json_str = re.sub(r'"$', '"}', json_str)
    
    # Handle truncated arrays
    if json_str.count('[') > json_str.count(']'):
        json_str += ']'
    
    # Handle truncated objects
    if json_str.count('{') > json_str.count('}'):
        json_str += '}'
    
    return json_str

def repair_truncated_json(json_str):
    """Attempt to repair truncated JSON from API responses"""
    print("Attempting to repair truncated JSON", file=sys.stderr)
    
    # Step 1: Identify which section is truncated
    if '"tools": [' in json_str and not re.search(r'"tools":\s*\[.*?\]', json_str, re.DOTALL):
        print("Detected truncation in tools array", file=sys.stderr)
        # Find where the tools array starts
        tools_start = json_str.find('"tools": [')
        if tools_start > 0:
            # Get the content before the tools array
            json_prefix = json_str[:tools_start] + '"tools": []'
            
            # Close any remaining open objects/arrays
            count_open_braces = json_prefix.count('{')
            count_close_braces = json_prefix.count('}')
            for _ in range(count_open_braces - count_close_braces):
                json_prefix += '}'
                
            try:
                # See if we can parse it now
                json.loads(json_prefix)
                print("Successfully repaired truncated JSON", file=sys.stderr)
                return json_prefix
            except json.JSONDecodeError as e:
                print(f"Failed to repair truncation: {e}", file=sys.stderr)
    
    # Try our general JSON cleaning function
    return clean_json_string(json_str)

def extract_name_from_text(text):
    """Extract a name from resume text using common patterns"""
    # First check for the exact format in the sample (Name followed by Email:)
    first_line_match = re.search(r'^([A-Za-z\s-]+)(?:\s+Email:)', text, re.MULTILINE)
    if first_line_match:
        return first_line_match.group(1).strip()
    
    # Pattern for name at beginning of resume or after common headers
    name_patterns = [
        r'^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s',  # Name at start
        r'(?:Name|PERSONAL\s+INFORMATION|PROFILE)[:;-]?\s*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
        r'^([A-Z][A-Z\s]+)\s',  # All caps name at beginning
    ]
    
    for pattern in name_patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            return match.group(1).strip()
    
    # Fallback: take first line that looks like a name (no special chars, reasonable length)
    lines = text.strip().split('\n')
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if 2 <= len(line.split()) <= 4 and re.match(r'^[A-Za-z\s.-]+$', line) and len(line) < 40:
            return line
    
    return "Unknown Name"

def extract_contact_info(text):
    """Extract contact information from resume text"""
    contact_info = {
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": ""
    }
    
    # Email extraction
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        contact_info["email"] = email_match.group(0)
    
    # Phone extraction - various formats
    phone_patterns = [
        r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # International
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US/CA
        r'\d{10,12}'  # Simple digits
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group(0)
            break
    
    # LinkedIn
    linkedin_match = re.search(r'linkedin\.com/in/[a-zA-Z0-9_-]+', text)
    if linkedin_match:
        contact_info["linkedin"] = "https://" + linkedin_match.group(0)
    elif "linkedin" in text.lower():
        # Try to find any LinkedIn reference
        linkedin_line = [line for line in text.split('\n') if 'linkedin' in line.lower()]
        if linkedin_line:
            contact_info["linkedin"] = linkedin_line[0].strip()
    
    # GitHub
    github_match = re.search(r'github\.com/[a-zA-Z0-9_-]+', text)
    if github_match:
        contact_info["github"] = "https://" + github_match.group(0)
    elif "github" in text.lower():
        github_line = [line for line in text.split('\n') if 'github' in line.lower()]
        if github_line:
            contact_info["github"] = github_line[0].strip()
    
    return contact_info

def extract_skills(text):
    """Extract skills from resume text"""
    skills = {
        "technical_skills": [],
        "soft_skills": [],
        "tools": []
    }
    
    # Common technical skills to look for
    tech_skills = ["Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Ruby", "PHP", "Swift", 
                  "React", "Angular", "Vue", "Node.js", "Express", "Django", "Flask", "Spring", 
                  "HTML", "CSS", "SQL", "NoSQL", "MongoDB", "PostgreSQL", "MySQL", "Oracle", 
                  "REST", "GraphQL", "AWS", "Azure", "GCP", "Docker", "Kubernetes", "CI/CD"]
    
    # Common tools
    tools = ["Git", "GitHub", "GitLab", "JIRA", "Confluence", "Jenkins", "Travis CI", "CircleCI",
             "VS Code", "Visual Studio", "IntelliJ", "PyCharm", "Eclipse", "Postman", "Figma",
             "Adobe", "Office", "Excel", "PowerPoint", "Word", "Outlook", "Teams", "Slack"]
    
    # Soft skills
    soft_skills = ["Communication", "Teamwork", "Problem Solving", "Time Management", "Leadership",
                  "Critical Thinking", "Adaptability", "Creativity", "Collaboration", "Agile", 
                  "Scrum", "Project Management", "Mentoring", "Presentation"]
    
    # Find skills section in resume
    skills_section = ""
    skill_headers = ["SKILLS", "TECHNICAL SKILLS", "TECHNOLOGIES", "COMPETENCIES", "EXPERTISE"]
    
    # Try to find a skills section
    lines = text.split('\n')
    in_skills_section = False
    
    for i, line in enumerate(lines):
        if any(header in line.upper() for header in skill_headers):
            in_skills_section = True
            skills_section = line + "\n"
        elif in_skills_section:
            if i+1 < len(lines) and any(header in lines[i+1].upper() for header in 
                                    ["EDUCATION", "EXPERIENCE", "WORK", "EMPLOYMENT", "PROJECTS"]):
                in_skills_section = False
            else:
                skills_section += line + "\n"
    
    # Extract from whole document if no dedicated section found
    text_to_search = skills_section if skills_section else text
    
    # Look for technical skills
    for skill in tech_skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_to_search, re.IGNORECASE):
            # Use the properly capitalized version from our list
            if skill not in skills["technical_skills"]:
                skills["technical_skills"].append(skill)
    
    # Look for tools
    for tool in tools:
        pattern = r'\b' + re.escape(tool) + r'\b'
        if re.search(pattern, text_to_search, re.IGNORECASE):
            if tool not in skills["tools"]:
                skills["tools"].append(tool)
    
    # Look for soft skills
    for skill in soft_skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_to_search, re.IGNORECASE):
            if skill not in skills["soft_skills"]:
                skills["soft_skills"].append(skill)
    
    return skills

def generate_fallback_json(text):
    """Generate basic JSON structure from resume text when the API fails"""
    print("Generating basic JSON using fallback extraction methods", file=sys.stderr)
    
    result = {
        "name": extract_name_from_text(text),
        "contact_info": extract_contact_info(text),
        "education": [
            {
                "institution": "",
                "degree": "",
                "date": "",
                "description": "Education details could not be automatically extracted"
            }
        ],
        "experience": [
            {
                "company": "",
                "position": "",
                "date": "",
                "description": "Experience details could not be automatically extracted"
            }
        ],
        "skill": extract_skills(text)
    }
    
    # Extract position/title if possible
    position_match = re.search(r'(?:^|\n)(.{0,50}(?:Developer|Engineer|Architect|Designer|Manager|Analyst|Consultant|Administrator).{0,30})(?:$|\n)', text)
    if position_match:
        position = position_match.group(1).strip()
        if len(position) < 60:  # Avoid overly long matches
            result["experience"][0]["position"] = position
    
    # Attempt to find any education reference
    edu_match = re.search(r'(?:University|College|Institute|School)[^\n.]{3,50}', text)
    if edu_match:
        result["education"][0]["institution"] = edu_match.group(0).strip()
    
    # Look for degree
    degree_match = re.search(r'(?:Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.|B\.E\.|M\.E\.)[^\n.]{3,50}', text)
    if degree_match:
        result["education"][0]["degree"] = degree_match.group(0).strip()
    
    # Try to find company name in context of work experience
    company_match = re.search(r'(?:at|for|with)\s+([A-Z][A-Za-z0-9\s&,.]+?)(?:\s+as|\s+from|\s+in|\s+\(|\s*\n)', text)
    if company_match:
        company = company_match.group(1).strip()
        if len(company) < 40:  # Avoid overly long matches
            result["experience"][0]["company"] = company
    
    # Convert to JSON string with proper indentation
    return json.dumps(result, indent=2)

def process_resume(file_path):
    """Main function to process resume file"""
    print(f"Processing resume: {file_path}", file=sys.stderr)
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} does not exist", file=sys.stderr)
        return json.dumps({"error": "File not found"})
    
    # Get file extension to handle different file types
    _, file_ext = os.path.splitext(file_path.lower())
    
    # Extract text based on file type
    if file_ext == '.pdf':
        text = extract_pdf_text(file_path)
    elif file_ext in ('.docx', '.doc'):
        text = extract_docx_text(file_path)
    else:
        # For other file types, attempt to read as text
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # If UTF-8 fails, try with latin-1 which should never fail
            try:
                print(f"UTF-8 decode failed, trying with latin-1 encoding", file=sys.stderr)
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
            except Exception as e:
                print(f"Error reading file {file_path}: {e}", file=sys.stderr)
                text = ""
    
    if not text:
        print(f"Error: Could not extract text from {file_path}", file=sys.stderr)
        return json.dumps({"error": "Failed to extract text from resume"})
    
    # Clean text
    cleaned_text = clean_text(text)
    print(f"Extracted {len(cleaned_text)} characters of text", file=sys.stderr)
    
    # Print sample for debugging
    print("\n--- Sample of Extracted Text ---", file=sys.stderr)
    print(cleaned_text[:500] + "..." if len(cleaned_text) > 500 else cleaned_text, file=sys.stderr)
    print("--- End Sample Text ---\n", file=sys.stderr)
    
    # Send text to DeepSeek AI for parsing
    json_text = send_to_deepseek_ai(cleaned_text)
    
    # If AI parsing failed, return error
    if not json_text:
        return json.dumps({"error": "Failed to parse resume with AI"})
    
    # Try to validate the JSON before returning
    try:
        parsed_json = json.loads(json_text)
        # Make sure it follows the expected structure
        if not isinstance(parsed_json, dict) or "name" not in parsed_json:
            print("Warning: AI returned JSON doesn't match expected structure", file=sys.stderr)
    except json.JSONDecodeError:
        print("Warning: AI returned invalid JSON", file=sys.stderr)
    
    # Return the JSON directly
    return json_text

if __name__ == "__main__":
    try:
        if len(sys.argv) < 2:
            print(json.dumps({'error': 'Please provide a file path'}))
            sys.exit(1)
        
        file_path = sys.argv[1]
        result = process_resume(file_path)
        
        # Just print the JSON directly
        if result:
            print(result)
        else:
            print(json.dumps({'error': 'Failed to parse resume'}))
            
    except Exception as e:
        print(json.dumps({'error': f'Error processing resume: {str(e)}'}))
        print(f"Error: {e}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
    
    sys.exit(0)